from docx import Document

# Load the input and output documents
input_doc = Document('ReferralFormFilled.docx')
output_doc = Document('OutputFormEmpty.docx')

# Function to extract text from paragraphs based on a keyword
def extract_text(doc, keyword):
    for para in doc.paragraphs:
        if keyword in para.text:
            return para.text.split(keyword)[1].strip()
    return ''  # Return an empty string if the keyword is not found

def split_name(name):
    name_parts = name.split(' ')
    return name_parts[0], ' '.join(name_parts[1:])

def split_disabilities(disabilities):
    return disabilities.split(',')

# Extract data from input document
data = {
    'Last name': split_name(extract_text(input_doc, "Name (First & Last):"))[1],
    'First name': split_name(extract_text(input_doc, "Name (First & Last):"))[0],
    'NHI number: ': extract_text(input_doc, "National Health Index (NHI) Number:"),
    'Title': extract_text(input_doc, "Title (Mr/Mrs/Ms/etc.):"),
    'Address': extract_text(input_doc, "Address:"),
    'Mobile': extract_text(input_doc, "Contact Number:"),
    'Email': extract_text(input_doc, "Email:"),
    'Date of Birth': extract_text(input_doc, "Date of Birth:"),
    'Gender': extract_text(input_doc, "Gender:"),
    'Ethnicity': extract_text(input_doc, "Ethnicity:"),
    'Iwi / Hapū (if Māori):' : extract_text(input_doc, "Iwi / Hapū (if Māori):"),
    'First language': extract_text(input_doc, "First Language (if not English):"),
    'Interpreter required': extract_text(input_doc, "Interpreter Required:"),
    'Cultural support required': extract_text(input_doc, "Cultural Support Required:"),
    "Communication needs": extract_text(input_doc, "Communication Needs:"),
    "Preferred contact method": extract_text(input_doc, "Preferred Contact Method:"),
    "Name of GP": extract_text(input_doc, "Doctor/GP Name:"),
    "GP's phone number": extract_text(input_doc, "Medical Centre Contact Number:"),
    "Primary disability": split_disabilities(extract_text(input_doc, "Disability Name/Type:"))[0],
}

print(data)

# Function to append data to placeholders in the output document
# def append_data_to_placeholders(doc, replacements):
#     for para in doc.paragraphs:
#         for key, value in replacements.items():
#             if key in para.text:
#                 # Append the extracted value to the existing text
#                 para.text += f' {value if value else ""}'

# # Append data to output document
# append_data_to_placeholders(output_doc, data)

# # Save the updated output document
# output_doc.save('OutputFormFilled.docx')
# print("Data has been appended successfully.")
