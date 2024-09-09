import whisper
import numpy as np
from pydub import AudioSegment
from pydub.silence import detect_silence
import noisereduce as nr
from docx import Document

audio = AudioSegment.from_file("Deacon.m4a")

normalized_audio = audio.apply_gain(-audio.dBFS)

audio_np = np.array(normalized_audio.get_array_of_samples())

reduced_noise_audio = nr.reduce_noise(y=audio_np, sr=audio.frame_rate)

reduced_noise_audio_segment = AudioSegment(
    reduced_noise_audio.tobytes(), 
    frame_rate=audio.frame_rate, 
    sample_width=audio.sample_width, 
    channels=audio.channels
)

silences = detect_silence(reduced_noise_audio_segment, min_silence_len=2000, silence_thresh=-35)

if silences:
    start_trim, end_trim = silences[0][0], silences[-1][1]
    trimmed_audio = reduced_noise_audio_segment[start_trim:end_trim]
else:
    trimmed_audio = reduced_noise_audio_segment

trimmed_audio.export("Deacon_trimmed.wav", format="wav")

model = whisper.load_model("large")
result = model.transcribe("Deacon_trimmed.wav", language="en", verbose=True)
print("Transcription of processed audio:")
document = Document()
document.add_heading('Transcription of processed audio:', level=1)
document.add_paragraph(result["text"])
document.save('Transcription-processed.docx')

result_original = model.transcribe("Deacon.m4a", language="en", verbose=True)
original_document = Document()
original_document.add_heading('Transcription of original audio:', level=1)
original_document.add_paragraph(result_original["text"])
original_document.save('Transcription-original.docx')