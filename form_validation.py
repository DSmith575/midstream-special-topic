from flask import Flask, request, jsonify
from docx import Document
from docx.shared import Pt
from datetime import datetime
import os

app = Flask(__name__)

# Constants
UPLOAD_FOLDER = './uploads'
OUTPUT_FOLDER = './outputs'

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

if not os.path.exists(OUTPUT_FOLDER):
    os.makedirs(OUTPUT_FOLDER)

# Function to extract text from paragraphs based on a keyword
def extract_text(doc, keyword, start_text=None, end_text=None):
    """Extract text from document paragraphs based on a keyword."""
    start_found = not bool(start_text)
    for para in doc.paragraphs:
        if start_text and start_text in para.text:
            start_found = True
            continue

        if end_text and end_text in para.text:
            start_found = False
            continue

        if start_found and keyword in para.text:
            return para.text.split(keyword)[1].strip()

def fill_fields(output_doc, data, start_text=None, end_text=None):
    """Fill fields in the output document with extracted data."""
    start_found = not bool(start_text)
    for paragraph in output_doc.paragraphs:
        if start_text and start_text in paragraph.text:
            start_found = True
            continue

        if end_text and end_text in paragraph.text:
            start_found = False
            continue

        if start_found:
            for key, value in data.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, f"{key} {value}")

def split_name(name):
    """Split a full name into first and last name."""
    name_parts = name.split(' ')
    return name_parts[0], ' '.join(name_parts[1:])

def split_disabilities(disabilities):
    """Split a comma-separated list of disabilities."""
    return disabilities.split(',')

def insert_bullet_points(output_doc, data, start_text=None, end_text=None):
    """Insert bullet points into the document based on the data provided."""
    start_found = not bool(start_text)

    for paragraph in output_doc.paragraphs:
        if start_text and start_text in paragraph.text:
            start_found = True
            continue

        if end_text and end_text in paragraph.text:
            start_found = False
            continue

        if start_found:
            for key, value in data.items():
                if key in paragraph.text:
                    paragraph.clear()
                    items = value.split(', ')
                    for item in items:
                        new_run = paragraph.add_run(f"•  {item.strip()}")
                        new_run.font.name = 'Times New Roman'
                        new_run.font.size = Pt(12)
                        new_run.font.bold = True
                        new_run.add_break()

@app.route('/process-document', methods=['POST'])
def process_document():
    """Endpoint to process the document."""
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    input_path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(input_path)

    # Load the input document
    input_doc = Document(input_path)
    output_doc = Document('OutputFormEmpty.docx')

    # Extracting details
    assessment_details = {
        'The assessment was completed on:': datetime.now().strftime("%d/%m/%Y"),
        'Name:': extract_text(input_doc, "Name (First & Last):"),
        'NHI:': extract_text(input_doc, "National Health Index (NHI) Number:"),
    }

    client_details = {
        'Last name:': split_name(extract_text(input_doc, "Name (First & Last):", "Section 1: Personal Information", "Section 2: Medical Information"))[1],
        'First name:': split_name(extract_text(input_doc, "Name (First & Last):", "Section 1: Personal Information", "Section 2: Medical Information"))[0],
        'NHI number:': extract_text(input_doc, "National Health Index (NHI) Number:", "Section 1: Personal Information", "Section 2: Medical Information"),
        'Title:': extract_text(input_doc, "Title (Mr/Mrs/Ms/etc.):", "Section 1: Personal Information", "Section 2: Medical Information"),
        'Marital status:': extract_text(input_doc, "Marital Status:", "Section 1: Personal Information", "Section 2: Medical Information") or "Single",
        'Address:': extract_text(input_doc, "Address:", "Section 1: Personal Information", "Section 2: Medical Information"),
        'Mobile:': extract_text(input_doc, "Contact Number:", "Section 1: Personal Information", "Section 2: Medical Information"),
        'Email:': extract_text(input_doc, "Email:", "Section 1: Personal Information", "Section 2: Medical Information"),
        'Date of birth:': extract_text(input_doc, "Date of Birth:", "Section 1: Personal Information", "Section 2: Medical Information"),
        'Gender:': extract_text(input_doc, "Gender:", "Section 1: Personal Information", "Section 2: Medical Information"),
        'Ethnicity:': extract_text(input_doc, "Ethnicity:", "Section 1: Personal Information", "Section 2: Medical Information"),
        'Iwi / Hapū (if Māori):': extract_text(input_doc, "Iwi / Hapū (if Māori):", "Section 1: Personal Information", "Section 2: Medical Information"),
        'First language:': extract_text(input_doc, "First Language (if not English):", "Section 1: Personal Information", "Section 2: Medical Information"),
        'Interpreter required:': extract_text(input_doc, "Interpreter Required:", "Section 1: Personal Information", "Section 2: Medical Information"),
        'Cultural support required:': extract_text(input_doc, "Cultural Support Required:", "Section 1: Personal Information", "Section 2: Medical Information"),
        "Communication needs:": extract_text(input_doc, "Communication Needs:", "Section 1: Personal Information", "Section 2: Medical Information"),
        "Preferred contact method:": extract_text(input_doc, "Preferred Contact Method:", "Section 1: Personal Information", "Section 2: Medical Information"),
        "Name of GP:": extract_text(input_doc, "Doctor/GP Name:", "Section 2: Medical Information", "Section 3: Disability Information"),
        "GP’s phone number:": extract_text(input_doc, "Medical Centre Contact Number:", "Section 2: Medical Information", "Section 3: Disability Information"),
        "Primary disability:": split_disabilities(extract_text(input_doc, "Disability Name/Type:", "Section 3: Disability Information", "Section 4: Additional Information"))[0],
        "Interim eligibility:": extract_text(input_doc, "Interim Eligibility:", "Section 1: Personal Information", "Section 2: Medical Information") or "No",
    }

    alt_contact = {
        'Last name:': split_name(extract_text(input_doc, "Name (First & Last):", "Section 5: Alternative Contact (Optional)", "Section 6: Referrer’s Contact Details"))[1],
        'First name': split_name(extract_text(input_doc, "Name (First & Last):", "Section 5: Alternative Contact (Optional)", "Section 6: Referrer’s Contact Details"))[0],
        'Address:': extract_text(input_doc, "Address:", "Section 5: Alternative Contact (Optional)", "Section 6: Referrer’s Contact Details") or "N/A",
        'Home phone:': extract_text(input_doc, "Contact Number:", "Section 5: Alternative Contact (Optional)", "Section 6: Referrer’s Contact Details") or "N/A",
        'Mobile:': extract_text(input_doc, "Contact Number:", "Section 5: Alternative Contact (Optional)", "Section 6: Referrer’s Contact Details") or "N/A",
        'Email:': extract_text(input_doc, "Email:", "Section 5: Alternative Contact (Optional)", "Section 6: Referrer’s Contact Details") or "N/A",
        'Relationship to client:': extract_text(input_doc, "Relationship to the Person Referred:", "Section 5: Alternative Contact (Optional)", "Section 6: Referrer’s Contact Details") or "N/A",
        'Date of birth:': extract_text(input_doc, "Date of Birth:", "Section 5: Alternative Contact (Optional)", "Section 6: Referrer’s Contact Details") or "N/A",
    }

    disability_diagnosis = {
        '[Details]': extract_text(input_doc, "Disability Name/Type:", "Section 3: Disability Information", "Section 4: Additional Information"),
    }

    # Fill the output document with all required data
    fill_fields(output_doc, assessment_details, "Assessment Information", "Client Details")
    fill_fields(output_doc, client_details, "Client Details", "Disability / Diagnosis Details")
    fill_fields(output_doc, alt_contact, "Alternative Contact Details", "Emergency Contact Details")
    insert_bullet_points(output_doc, disability_diagnosis, "Disability / Diagnosis Details", "Reason for Assessment / Referral to Taikura Trust")

    date_created = datetime.now().strftime("%d-%m-%Y")
    output_path = os.path.join(OUTPUT_FOLDER, f"{date_created}_OutputForm.docx")
    output_doc.save(output_path)

    return jsonify({"message": "Data has been transferred and appended successfully.", "output_file": output_path}), 200

if __name__ == '__main__':
    app.run(debug=True)
