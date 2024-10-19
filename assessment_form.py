from openai import OpenAI
import pymupdf
import os
import time
from docx2pdf import convert
from docx import Document
from dotenv import load_dotenv

load_dotenv()

api_key = os.getenv('OPENAI_API_KEY')

PROCESSED_DIR = 'processed'
UPLOADS_DIR = 'uploads'


client = OpenAI(api_key=api_key)

def convert_to_pdf(docx_path):
    """Convert DOCX to PDF."""
    if not os.path.isfile(docx_path):
        print(f"{docx_path} does not exist.")
        raise FileNotFoundError(f"{docx_path} does not exist.")

def convert_docx_to_pdf(docx_path):
    try:
        pdf_path = docx_path.replace(".docx", ".pdf")
        convert(docx_path, pdf_path)  # Converts the DOCX file to PDF
        return pdf_path
    except Exception as e:
        print(f"Error converting {docx_path} to PDF: {str(e)}")
        raise

def save_form_data_to_doc(form_data):
    doc = Document()
    doc.add_heading('Processed Form Data', level=1)

    # Write each key-value pair in form_data to the document
    for item, response in form_data.items():
        doc.add_heading(item, level=2)
        doc.add_paragraph(str(response) if response else "No relevant information found.")

    output_filename = 'processed_data.docx'
    output_filepath = os.path.join(UPLOADS_DIR, output_filename)

    try:
        # Save the document to the specified path
        doc.save(output_filepath)
        return output_filepath
    except Exception as e:
        print(f"Failed to write processed data: {e}")
        return None


def extract_text_from_pdf(filepath):
    if not os.path.isfile(filepath):
        print(f"File does not exist: {filepath}")
        return None

    try:
        # Open the PDF file
        # open with pymu
        with pymupdf.open(filepath) as doc:
            text = ""
            # Extract text from each page
            for page in doc:
                text += page.get_text("text")
            
            return text

        with pymupdf.open(filepath) as doc:  # Using a context manager
            text = ""
            # Extract text from each page
            for page in doc:
                text += page.get_text("text")  # Extract text as plain text

            return text

    except Exception as e:
        print(f"Error extracting text from PDF: {e} - File: {filepath}")
        return None


def get_relevant_information(section, text):
    prompt = (f"Based on the following text, does the person mentioned have any issues related to {section}? "
              f"If yes, provide details. Text: {text}")
    
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt}
            ],
        )
        return response.choices[0].message.content.strip()

    except Exception as e:
        print(f"Error getting response from OpenAI: {e}")
        return None


def analyze_pdf_for_form(text):
    form_data = {}

    print("Analyzing PDF for form data...")
    sections = {
        'Household Management': [
            'Faecal smearing', 'Administering personal finances', 'Garden / lawns', 
            'Home safety', 'Laundry', 'Operating home heating appliances', 
            'Meal preparation', 'Shopping for necessary items', 'Other housework'
        ],
        'Self-care': [
            'Bathing, showering, washing self', 'Bed mobility', 'Dressing and / or undressing', 
            'Eating and drinking', 'Faecal smearing', 'Grooming and caring for body parts', 
            'Managing / preventing health problems', 'Managing medication', 
            'Menstrual management', 'Night Care', 'Night settling', 'Toileting'
        ],
        # Add other sections accordingly
    }

    print("Sections:", sections)  

    for section, items in sections.items():
        print(f"Analyzing section: {section}")
        for item in items:
            print(f"Looking for item: {item}") 
            response = get_relevant_information(item, text)
            
            if response is None:
                print(f"Warning: No relevant information found for {item}")
            else:
                print(f"Found relevant information for {item}: {response}")
            
            form_data[item] = response


    return form_data




def process_data(pdf_path):
    try:
        start_time = time.time()
        extracted_text = extract_text_from_pdf(pdf_path)
        form_data = analyze_pdf_for_form(extracted_text)
        save_doc = save_form_data_to_doc(form_data)
        pdf_path = convert_to_pdf(save_doc)

        print(f"Processed data saved as PDF: {pdf_path}")

        print(f"Processing time: {time.time() - start_time:.2f} seconds")

        return pdf_path
    except Exception as e:
        print(f"Error processing PDF: {str(e)}")
        raise


