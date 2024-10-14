import os
import sys
import numpy as np
from multiprocessing import Pool
from pydub import AudioSegment
from pydub.silence import detect_silence
from pydub.utils import make_chunks
import noisereduce as nr
import whisper
from docx import Document
import torch
import time

import logging
import traceback


logging.basicConfig(level=logging.ERROR, filename='error.log', filemode='a', 
                    format='%(asctime)s - %(levelname)s - %(message)s')


CHUNK_LENGTH_MS = 120000  # 2 minutes
SILENCE_THRESHOLD = -30    # in dBFS
MIN_SILENCE_LEN = 1000     # in milliseconds
PROCESSED_DIR = 'processed'
UPLOADS_DIR = 'uploads'

def create_directories():
    # Create necessary directories for processed audio and uploads.
    os.makedirs(PROCESSED_DIR, exist_ok=True)
    os.makedirs(UPLOADS_DIR, exist_ok=True)

def clean_existing_files(filename):
    # Remove existing files if they already exist.
    for file_path in [
        os.path.join(PROCESSED_DIR, f"{filename}.wav"),
        os.path.join(PROCESSED_DIR, f"{filename}-original.docx"),
        os.path.join(PROCESSED_DIR, f"{filename}-processed.docx")
    ]:
        if os.path.exists(file_path):
            os.remove(file_path)

def process_chunk(chunk):
    # Reduce noise in audio chunk and remove silence.
    chunk_np = np.array(chunk.get_array_of_samples())
    reduced_chunk = nr.reduce_noise(y=chunk_np, sr=chunk.frame_rate, prop_decrease=0.8)
    
    reduced_chunk_audio = AudioSegment(
        reduced_chunk.tobytes(),
        frame_rate=chunk.frame_rate,
        sample_width=chunk.sample_width,
        channels=chunk.channels
    )

    silences = detect_silence(reduced_chunk_audio, min_silence_len=MIN_SILENCE_LEN, silence_thresh=SILENCE_THRESHOLD)
    
    if silences:
        start_trim, end_trim = silences[0][0], silences[-1][1]
        return reduced_chunk_audio[start_trim:end_trim]
    return reduced_chunk_audio

def transcribe_audio(audio_path, model):
    # Transcribe audio using the Whisper model.
    return model.transcribe(audio_path, language="en", verbose=True)

def save_transcription_to_docx(transcription, filename, original=False):
    # Save transcription to a .docx file.
    document = Document()
    title = 'Transcription of original audio:' if original else 'Transcription of processed audio:'
    document.add_heading(title, level=1)
    document.add_paragraph(transcription["text"])
    docx_filename = f"{filename}-original.docx" if original else f"{filename}-processed.docx"
    document.save(os.path.join(PROCESSED_DIR, docx_filename))
    
    return os.path.join(PROCESSED_DIR, docx_filename)



def process_audio(audio_path):
    try:
        start_time = time.time()
        print("Starting transcription...")

        create_directories()

        audio = AudioSegment.from_file(audio_path)

        filename = os.path.splitext(os.path.basename(audio_path))[0]
        clean_existing_files(filename)

        chunks = make_chunks(audio, CHUNK_LENGTH_MS)

        # Process the audio chunks
        processed_chunks = [process_chunk(chunk) for chunk in chunks]

        combined_audio = sum(processed_chunks)
        
        processed_wav_path = os.path.join(PROCESSED_DIR, f"{filename}.wav")
        combined_audio.export(processed_wav_path, format="wav")

        device = "cuda" if torch.cuda.is_available() else "cpu"
        # device = "cuda"
        # if not torch.cuda.is_available():
        #   raise RuntimeError("CUDA is not available. Please check your installation and GPU.")
        model = whisper.load_model("large", device=device)

        # Transcribe the processed audio
        print("Transcribing processed audio...")
        processed_result = transcribe_audio(processed_wav_path, model)
        processed_document_path = save_transcription_to_docx(processed_result, filename)

        # Transcribe the original audio
        # print("Transcribing original audio...")
        # original_result = transcribe_audio(audio_path, model)
        # save_transcription_to_docx(original_result, filename, original=True)

        print('Finished Processing')
        print(f"Time taken: {time.time() - start_time} seconds")
        
        # Return both the processed audio path and the transcription result
        return processed_document_path
    except Exception as e:
        logging.error(f"An error occurred: {str(e)}")
        raise  # Re-raise the exception for the caller to handle



if __name__ == '__main__':
    if len(sys.argv) != 2:
        print("Usage: python audio_process.py <audio_path>")
        sys.exit(1)
    audio_path = sys.argv[1]
    process_audio(audio_path)
